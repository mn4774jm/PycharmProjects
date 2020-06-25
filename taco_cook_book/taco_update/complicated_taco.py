import requests
import docx
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt

# creating a variable for the url that generates random taco recipes
url = 'https://taco-1150.herokuapp.com/random/?full_taco=true'

# create lists and dictionary
taco_json = []
name_list = ['recipe1', 'recipe2', 'recipe3']
taco_dict = {}
taco_titles = []

# loop through and create a list containing json data from the 3 requests
for i in range(3):
    taco_json.append(requests.get(url).json())

# loop to create titles for each recipe
for i in range(3):
    taco_titles.append(taco_json[i]['base_layer']['name'] + ' with ' + taco_json[i]['mixin']['name']
                       + ' and ' + taco_json[i]['condiment']['name'])

# loop to grab individual items from the json data, store to list, and store to dictionary
for i in range(3):
    ingredient_list = [taco_json[i]['seasoning']['recipe'], taco_json[i]['condiment']['recipe'],
                       taco_json[i]['mixin']['recipe'], taco_json[i]['base_layer']['recipe'],
                       taco_json[i]['shell']['recipe']]
    taco_dict[name_list[i]] = ingredient_list

image = Image.open('tacopic.jpg')
image = image.resize((400,400))
image.save('tacopic.jpg')

# opening the taco pic again, using the ImageDraw to write on the picture, setting the font, font size, color, and location of the text, then saving the edited pic
image = Image.open('tacopic.jpg')
imgDraw = ImageDraw.Draw(image)
font = ImageFont.truetype('DejaVuSans.ttf', 20)
imgDraw.text([0, 350], 'Random Taco Cookbook', fill='black', font=font)
image.save('tacopic2.jpg')

# creating a new word document using the docx library
document = docx.Document()
# adding a heading (title), on the first line, and setting the font size
newHeading = document.add_heading(level=0)
writeHeading = newHeading.add_run('Random Taco Cookbook')
writeHeading.font.size = Pt(24)

# adding the edited taco picture to the title page
document.add_picture('tacopic2.jpg')

# adding a paragraph, with the Heading 1 format to make stand out. Then adding another paragraph with credit info
document.add_paragraph('Credits:\n', 'Heading 1')
document.add_paragraph(f'Taco Image: Photo by Tai\'s Captures on Unsplash\n'
                       'Taco Recipes From: https://taco-1150.herokuapp.com/random/?full_taco=true\n')

# creating a page break
document.add_page_break()

# adding a new heading (title) for the new page on the first line
newHeading = document.add_heading(level=0)
# new heading using the variable for the recipe 1 title string, setting the font size
writeHeading = newHeading.add_run(f'{taco_titles[0]}')

#need to sort out how to turn this into a loop. Program starts on line 78 of other program
writeHeading.font.size = Pt(24)
# adding new paragraphs of each recipe name, formatted to Heading 1 so that they stand out, pulling info from index
document.add_paragraph(f'{taco_json[0]["seasoning"]["name"]}\n', 'Heading 1')